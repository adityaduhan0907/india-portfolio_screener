import requests
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from bs4 import BeautifulSoup
finnhub_api_key = "d15ij81r01qhqto62tf0d15ij81r01qhqto62tfg"
newsapi_key = "75d648980fb542268b5a5255f6fc65e4"
companies = ["ASIANPAINT.NS", "FINCABLES.NS", "MPHASIS.NS"]
company_names = {
    "ASIANPAINT.NS": "ASIAN PAINTS",
    "FINCABLES.NS": "FINOLEX CABLES",
    "MPHASIS.NS": "MPHASIS"
}
end_date = datetime.strptime("2025-06-01", "%Y-%m-%d")
start_date = end_date - timedelta(days=5 * 365)
excel_path = "SWOT_Raw_Data.xlsx"
word_path = "SWOT_Analysis_Report.docx"
writer = pd.ExcelWriter(excel_path, engine='openpyxl')
document = Document()
document.add_heading('SWOT Analysis Report', 0)
def fetch_finnhub_financials(symbol):
    url = f"https://finnhub.io/api/v1/stock/financials-reported?symbol={symbol}&token={finnhub_api_key}"
    res = requests.get(url)
    return res.json().get('data', [])
def fetch_finnhub_ratios(symbol):
    url = f"https://finnhub.io/api/v1/stock/metric?symbol={symbol}&metric=all&token={finnhub_api_key}"
    res = requests.get(url)
    return res.json().get("metric", {})
def fetch_news_articles(symbol):
    url = f"https://newsapi.org/v2/everything?q={symbol}&language=en&sortBy=publishedAt&pageSize=5&apiKey={newsapi_key}"
    res = requests.get(url)
    return res.json().get("articles", [])
def scrape_nse_management_commentary(company_name):
    search_name = company_name.replace(" ", "%20")
    url = f"https://www.nseindia.com/api/search/autocomplete?q={search_name}"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        res = requests.get(url, headers=headers, timeout=10)
        if res.ok:
            json_data = res.json()
            if json_data and 'symbols' in json_data and json_data['symbols']:
                symbol_info = json_data['symbols'][0]
                code = symbol_info['symbol']
                company_url = f"https://www.nseindia.com/get-quotes/equity?symbol={code}"
                return f"Check NSE Commentary Page: {company_url}"
    except:
        pass
    return "NSE commentary not found."
def generate_swot_from_data(financials, ratios, news, commentary):
    strengths, weaknesses, opportunities, threats = [], [], [], []
    if ratios.get("roe") and ratios["roe"] > 15:
        strengths.append("High Return on Equity (ROE).")
    if ratios.get("debtToEquity") and ratios["debtToEquity"] > 2:
        weaknesses.append("High Debt to Equity Ratio.")

    for article in news:
        desc = article.get("description", "").lower()
        if "growth" in desc or "expansion" in desc:
            opportunities.append("Expansion-related developments.")
        if "risk" in desc or "fraud" in desc or "layoff" in desc:
            threats.append("Negative news reported.")

    if "not found" not in commentary:
        opportunities.append("New insights available from management commentary.")

    return strengths, weaknesses, opportunities, threats
for symbol in companies:
    name = company_names[symbol]
    print(f"Processing: {name}")

    fin = fetch_finnhub_financials(symbol)
    ratios = fetch_finnhub_ratios(symbol)
    news = fetch_news_articles(symbol)
    commentary = scrape_nse_management_commentary(name)
    fin_df = pd.json_normalize(fin)
    fin_df.to_excel(writer, sheet_name=f"{symbol}_financials", index=False)

    ratios_df = pd.DataFrame([ratios])
    ratios_df.to_excel(writer, sheet_name=f"{symbol}_ratios", index=False)
    s, w, o, t = generate_swot_from_data(fin, ratios, news, commentary)
    document.add_heading(f"{name}", level=1)
    document.add_paragraph("NSE Commentary URL or Status: " + commentary)

    document.add_heading("Strengths", level=2)
    for item in s: document.add_paragraph(item, style='List Bullet')

    document.add_heading("Weaknesses", level=2)
    for item in w: document.add_paragraph(item, style='List Bullet')

    document.add_heading("Opportunities", level=2)
    for item in o: document.add_paragraph(item, style='List Bullet')

    document.add_heading("Threats", level=2)
    for item in t: document.add_paragraph(item, style='List Bullet')
writer.close()
document.save(word_path)

print("SWOT analysis completed!")